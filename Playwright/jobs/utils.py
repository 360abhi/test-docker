from playwright.sync_api import Page
from docx import Document
from docx.shared import Inches
import os

class Utils:

    def  __init__(self,page:Page,logger):
        self.page = page
        self.logger = logger

    def append_to_docx(self,path:str):
        image_name = f"{path}.png"
        doc_path = f"{path}.docx"

        self.page.screenshot(path=image_name)
        
        if os.path.exists(doc_path):
            doc = Document(doc_path)
        else:
            doc = Document()

        doc.add_paragraph("Screenshot Added")
        doc.add_picture(image_name,width=Inches(6))

        doc.save(doc_path)
        os.remove(image_name)

        self.logger.info(f"Screenshot appended to {doc_path}")